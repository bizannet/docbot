# documents/formats/word.py
import tempfile
from pathlib import Path
from typing import Union

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup


def _add_paragraph(doc: Document, text: str, bold: bool = False, align: str = 'left'):
    """Вспомогательная функция для добавления параграфа"""
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)


def generate_word(
        html_content: str,
        output_filename: str,
        output_dir: Union[str, Path, None] = None
) -> str:
    """
    Конвертирует HTML-строку в DOCX-файл.
    Упрощённая версия для совместимости.
    """
    if output_dir is None:
        output_dir = tempfile.gettempdir()

    output_path = Path(output_dir) / output_filename

    doc = Document()

    # Настраиваем базовый стиль
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Парсим HTML
    soup = BeautifulSoup(html_content, 'html.parser')

    # Удаляем лишнее
    for tag in soup(['script', 'style', 'head', 'meta']):
        tag.decompose()

    # Проходим по элементам
    for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'div', 'li']):
        text = element.get_text(strip=True)
        if not text:
            continue

        tag_name = element.name

        if tag_name in ['h1', 'h2']:
            _add_paragraph(doc, text, bold=True, align='center')
        elif tag_name == 'h3':
            _add_paragraph(doc, text, bold=True)
        elif tag_name == 'li':
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(text)
        else:
            _add_paragraph(doc, text, align='justify')

    doc.save(str(output_path))
    return str(output_path)


def generate_word_from_file(
        template_path: Union[str, Path],
        context: dict,
        output_filename: str,
        output_dir: Union[str, Path, None] = None
) -> str:
    """
    Рендерит шаблон + конвертирует в Word (удобный wrapper).
    """
    from documents.base import DocumentGenerator

    gen = DocumentGenerator()
    html_content = gen.render_template(template_path, context)

    return generate_word(html_content, output_filename, output_dir)